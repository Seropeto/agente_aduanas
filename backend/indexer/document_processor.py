"""
Procesador de documentos para el sistema RAG de Aduanas Chile.
Soporta: PDF (pdfplumber + OCR fallback), DOCX, TXT.
Genera chunks de ~500 palabras con solapamiento de ~50 palabras.
"""
import hashlib
import io
import logging
import re
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)

# Umbral para activar OCR: si el texto extraído tiene menos de 100 chars
OCR_THRESHOLD = 100
CHUNK_SIZE = 500   # palabras
CHUNK_OVERLAP = 50  # palabras


class DocumentProcessor:
    """
    Procesa documentos de diferentes formatos y los divide en chunks
    para indexación en el vector store.
    """

    def process_file(self, file_path: str | Path, metadata: dict[str, Any] | None = None) -> list[dict[str, Any]]:
        """
        Procesa un archivo y retorna una lista de chunks con metadata.

        Args:
            file_path: Ruta al archivo a procesar.
            metadata: Metadata adicional a incluir en los chunks.

        Returns:
            Lista de dicts con 'text', 'metadata' y 'chunk_index'.
        """
        file_path = Path(file_path)
        if not file_path.exists():
            logger.error(f"Archivo no encontrado: {file_path}")
            return []

        suffix = file_path.suffix.lower()
        metadata = metadata or {}

        # Añadir metadata del archivo
        metadata.setdefault("filename", file_path.name)
        metadata.setdefault("file_path", str(file_path))
        metadata.setdefault("doc_id", self._compute_file_hash(file_path))

        try:
            if suffix == ".pdf":
                text = self._extract_pdf(file_path)
            elif suffix in (".docx", ".doc"):
                text = self._extract_docx(file_path)
            elif suffix == ".txt":
                text = self._extract_txt(file_path)
            elif suffix in (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"):
                text = self._extract_image(file_path)
            else:
                logger.warning(f"Formato no soportado: {suffix}")
                return []
        except Exception as e:
            logger.error(f"Error procesando {file_path}: {e}")
            return []

        if not text or len(text.strip()) < 10:
            if suffix in (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"):
                title = metadata.get("title", file_path.stem)
                source = metadata.get("source", "")
                date = metadata.get("date", "")
                content_type = metadata.get("content_type", "imagen")
                placeholder = (
                    f"Documento: {title}. "
                    f"Tipo: {content_type}. "
                    f"{'Fecha: ' + date + '. ' if date else ''}"
                    f"{'Origen: ' + source + '. ' if source else ''}"
                    f"Nombre de archivo: {file_path.name}. "
                    f"[Imagen sin texto detectable por OCR]"
                )
                logger.warning(f"No se detectó texto en imagen {file_path.name}. Indexando con metadata.")
                chunks = [{"text": placeholder, "metadata": {**metadata, "chunk_index": 0, "total_chunks": 1}}]
                return chunks
            logger.warning(f"No se pudo extraer texto de {file_path}")
            return []

        chunks = self._chunk_text(text)
        result = []
        for i, chunk in enumerate(chunks):
            chunk_meta = {**metadata, "chunk_index": i, "total_chunks": len(chunks)}
            result.append({"text": chunk, "metadata": chunk_meta})

        logger.info(f"Procesado {file_path.name}: {len(result)} chunks generados")
        return result

    def process_text(self, text: str, metadata: dict[str, Any] | None = None) -> list[dict[str, Any]]:
        """
        Procesa texto directo (por ejemplo, de scrapers web) y lo divide en chunks.

        Args:
            text: Texto a procesar.
            metadata: Metadata de la fuente.

        Returns:
            Lista de chunks con metadata.
        """
        if not text or len(text.strip()) < 10:
            return []

        metadata = metadata or {}
        # Generar doc_id si no existe
        if "doc_id" not in metadata:
            metadata["doc_id"] = hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]

        chunks = self._chunk_text(text)
        result = []
        for i, chunk in enumerate(chunks):
            chunk_meta = {**metadata, "chunk_index": i, "total_chunks": len(chunks)}
            result.append({"text": chunk, "metadata": chunk_meta})

        return result

    def compute_text_hash(self, text: str) -> str:
        """Calcula el hash SHA-256 de un texto."""
        return hashlib.sha256(text.encode("utf-8")).hexdigest()

    def compute_file_hash(self, file_path: str | Path) -> str:
        """Calcula el hash SHA-256 del contenido de un archivo."""
        return self._compute_file_hash(Path(file_path))

    # ------------------------------------------------------------------ #
    # Extracción de texto                                                   #
    # ------------------------------------------------------------------ #

    def _extract_pdf(self, file_path: Path) -> str:
        """
        Extrae texto de un PDF usando pdfplumber.
        Si el resultado tiene menos de OCR_THRESHOLD chars, activa OCR con pytesseract.
        """
        text = self._extract_pdf_pdfplumber(file_path)

        if len(text.strip()) < OCR_THRESHOLD:
            logger.info(
                f"PDF con poco texto ({len(text)} chars), activando OCR: {file_path.name}"
            )
            ocr_text = self._extract_pdf_ocr(file_path)
            if ocr_text and len(ocr_text.strip()) > len(text.strip()):
                text = ocr_text

        return text

    def _extract_pdf_pdfplumber(self, file_path: Path) -> str:
        """Extrae texto de PDF usando pdfplumber."""
        try:
            import pdfplumber
            pages_text = []
            with pdfplumber.open(str(file_path)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
            return "\n\n".join(pages_text)
        except ImportError:
            logger.warning("pdfplumber no disponible, intentando PyPDF2")
            return self._extract_pdf_pypdf2(file_path)
        except Exception as e:
            logger.warning(f"Error con pdfplumber en {file_path.name}: {e}")
            return self._extract_pdf_pypdf2(file_path)

    def _extract_pdf_pypdf2(self, file_path: Path) -> str:
        """Extrae texto de PDF usando PyPDF2 como fallback."""
        try:
            import PyPDF2
            pages_text = []
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
            return "\n\n".join(pages_text)
        except Exception as e:
            logger.error(f"Error con PyPDF2 en {file_path.name}: {e}")
            return ""

    def _extract_pdf_ocr(self, file_path: Path) -> str:
        """
        Convierte páginas del PDF a imágenes y aplica OCR con pytesseract.
        Solo se activa cuando pdfplumber no extrae suficiente texto.
        """
        try:
            import pytesseract
            from PIL import Image
            import pdfplumber

            pages_text = []
            with pdfplumber.open(str(file_path)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        # Convertir página a imagen
                        img = page.to_image(resolution=200)
                        pil_img = img.original

                        # OCR en español e inglés
                        ocr_text = pytesseract.image_to_string(
                            pil_img, lang="spa+eng",
                            config="--psm 1 --oem 3"
                        )
                        if ocr_text.strip():
                            pages_text.append(ocr_text)
                    except Exception as e:
                        logger.warning(f"Error OCR en página {page_num}: {e}")

            return "\n\n".join(pages_text)
        except ImportError:
            logger.warning("pytesseract o dependencias OCR no disponibles")
            return ""
        except Exception as e:
            logger.error(f"Error en OCR para {file_path.name}: {e}")
            return ""

    def _extract_docx(self, file_path: Path) -> str:
        """Extrae texto de un archivo DOCX."""
        try:
            from docx import Document
            doc = Document(str(file_path))
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            # También extraer tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error extrayendo DOCX {file_path.name}: {e}")
            return ""

    def _extract_image(self, file_path: Path) -> str:
        """Extrae texto de una imagen usando OCR con pytesseract."""
        try:
            import pytesseract
            from PIL import Image

            # Configurar Tesseract en Windows
            import sys, os, shutil
            if sys.platform == "win32":
                if not shutil.which("tesseract"):
                    for candidate in [
                        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                        r"C:\Users\HP\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
                    ]:
                        if Path(candidate).exists():
                            pytesseract.pytesseract.tesseract_cmd = candidate
                            break
                # Usar tessdata del usuario si el de Program Files no tiene spa
                user_tessdata = Path.home() / "tessdata"
                sys_tessdata = Path(r"C:\Program Files\Tesseract-OCR\tessdata")
                if user_tessdata.exists() and (user_tessdata / "spa.traineddata").exists():
                    if not (sys_tessdata / "spa.traineddata").exists():
                        os.environ["TESSDATA_PREFIX"] = str(user_tessdata)

            img = Image.open(str(file_path))
            # Convertir a RGB si es necesario (TIFF, BMP pueden ser otros modos)
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")

            # Intentar con español+inglés, si falla usar solo inglés
            try:
                text = pytesseract.image_to_string(img, lang="spa+eng", config="--psm 1 --oem 3")
            except Exception:
                text = pytesseract.image_to_string(img, lang="eng", config="--psm 1 --oem 3")
            logger.info(f"OCR completado en {file_path.name}: {len(text)} chars extraídos")
            return text
        except ImportError:
            logger.error("pytesseract o Pillow no disponibles para procesar imagen")
            return ""
        except Exception as e:
            logger.error(f"Error extrayendo texto de imagen {file_path.name}: {e}")
            return ""

    def _extract_txt(self, file_path: Path) -> str:
        """Extrae texto de un archivo TXT."""
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
        for enc in encodings:
            try:
                return file_path.read_text(encoding=enc)
            except (UnicodeDecodeError, Exception):
                continue
        logger.error(f"No se pudo leer {file_path.name} con ninguna codificación")
        return ""

    # ------------------------------------------------------------------ #
    # Chunking                                                              #
    # ------------------------------------------------------------------ #

    def _chunk_text(self, text: str) -> list[str]:
        """
        Divide el texto en chunks de CHUNK_SIZE palabras con CHUNK_OVERLAP de solapamiento.
        Intenta respetar límites de párrafos cuando es posible.
        """
        if not text:
            return []

        # Limpiar el texto
        text = self._normalize_text(text)

        # Tokenizar por palabras (conservando espacios implícitamente)
        words = text.split()

        if len(words) <= CHUNK_SIZE:
            return [text]

        chunks = []
        start = 0

        while start < len(words):
            end = min(start + CHUNK_SIZE, len(words))
            chunk_words = words[start:end]
            chunk_text = " ".join(chunk_words)

            # Intentar terminar en punto final para chunks no del último
            if end < len(words):
                last_50 = " ".join(chunk_words[-50:])
                last_period = last_50.rfind(".")
                if last_period > 0:
                    words_after_period = len(last_50[last_period + 1:].split())
                    adjusted_end = end - words_after_period
                    # Solo ajustar si no retrocedemos más allá del inicio del chunk
                    if adjusted_end > start:
                        chunk_words = words[start:adjusted_end]
                        chunk_text = " ".join(chunk_words)
                        end = adjusted_end

            if chunk_text.strip():
                chunks.append(chunk_text.strip())

            # Avanzar siempre hacia adelante
            next_start = end - CHUNK_OVERLAP
            start = next_start if next_start > start else end

        return chunks

    def _normalize_text(self, text: str) -> str:
        """Normaliza el texto eliminando caracteres problemáticos."""
        # Reemplazar múltiples espacios en blanco por uno solo
        text = re.sub(r"[ \t]+", " ", text)
        # Reemplazar múltiples saltos de línea
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Eliminar caracteres de control excepto newlines
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        return text.strip()

    def _compute_file_hash(self, file_path: Path) -> str:
        """Calcula el hash SHA-256 del contenido del archivo."""
        try:
            sha256 = hashlib.sha256()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(65536), b""):
                    sha256.update(chunk)
            return sha256.hexdigest()[:16]
        except Exception as e:
            logger.error(f"Error calculando hash de {file_path}: {e}")
            return hashlib.sha256(str(file_path).encode()).hexdigest()[:16]
